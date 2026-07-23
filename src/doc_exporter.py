# -*- coding: utf-8 -*-
"""
Markdown → Word / Excel 导出器
- Word: 解析 Markdown 标题/表格/代码块/列表/粗体，支持中文字体回退
- Excel: FMEA 表格专用导出，带条件着色和措施跟踪 Sheet
"""

import io
import re
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ======================================================================
# Word 导出
# ======================================================================

def export_to_word(title: str, markdown: str) -> bytes:
    """
    将 Markdown 内容导出为 Word (.docx) 文件，返回文件字节流。

    Args:
        title:    文档标题
        markdown: Markdown 格式文本

    Returns:
        .docx 文件的字节内容
    """
    doc = Document()

    # 设置默认字体与中文字体回退
    _setup_default_font(doc)

    # 添加文档标题
    heading = doc.add_heading(title, level=0)
    _set_chinese_font(heading.runs[0] if heading.runs else None)

    # 逐行解析 Markdown 并写入文档
    elements = _parse_markdown_elements(markdown)
    for elem in elements:
        _write_element_to_doc(doc, elem)

    # 添加 AI 免责声明
    _add_disclaimer(doc)

    # 保存到内存缓冲区
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ======================================================================
# Excel (FMEA) 导出
# ======================================================================

def export_fmea_to_excel(markdown: str) -> bytes:
    """
    从 FMEA 的 Markdown 文本中提取表格，导出为格式化的 Excel 文件。

    特性：
    - 表头蓝底白字
    - RPN ≥ 200 红色着色，RPN ≥ 100 黄色着色
    - 自动列宽
    - 冻结首行
    - 额外创建「措施跟踪」Sheet

    Args:
        markdown: FMEA 的 Markdown 文本

    Returns:
        .xlsx 文件的字节内容
    """
    wb = Workbook()

    # 提取所有 Markdown 表格
    tables = _extract_markdown_tables(markdown)

    if not tables:
        # 没有表格时创建一个空表
        ws = wb.active
        ws.title = "FMEA分析"
        ws["A1"] = "未检测到 FMEA 表格数据"
    else:
        # 第一个表格放入主 Sheet
        ws = wb.active
        ws.title = "FMEA分析"
        _write_table_to_sheet(ws, tables[0])

        # 如果有多个表格（按维度拆分），追加到后续 Sheet
        for idx, table in enumerate(tables[1:], start=2):
            extra_ws = wb.create_sheet(title=f"FMEA-{idx}")
            _write_table_to_sheet(extra_ws, table)

    # 创建「措施跟踪」Sheet
    tracking_ws = wb.create_sheet(title="措施跟踪")
    _build_tracking_sheet(tracking_ws, tables)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ======================================================================
# 内部辅助函数 —— Markdown 解析
# ======================================================================

# Markdown 元素类型常量
_ELEM_HEADING = "heading"
_ELEM_TABLE = "table"
_ELEM_CODE_BLOCK = "code_block"
_ELEM_LIST_ITEM = "list_item"
_ELEM_PARAGRAPH = "paragraph"

# 正则模式
_RE_HEADING = re.compile(r'^(#{1,6})\s+(.+)$')
_RE_TABLE_SEP = re.compile(r'^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$')
_RE_TABLE_ROW = re.compile(r'^\|(.+)\|?\s*$')
_RE_CODE_FENCE = re.compile(r'^(`{3,}|~{3,})(\w*)\s*$')
_RE_LIST_ITEM = re.compile(r'^(\s*)([-*+]|\d+[.)]) (.+)$')
_RE_BOLD = re.compile(r'\*\*(.+?)\*\*')


def _parse_markdown_elements(md: str) -> list:
    """将 Markdown 文本解析为结构化元素列表。"""
    lines = md.split('\n')
    elements = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # 代码块
        fence_match = _RE_CODE_FENCE.match(line)
        if fence_match:
            fence_char = fence_match.group(1)
            lang = fence_match.group(2)
            code_lines = []
            i += 1
            while i < n:
                if lines[i].strip().startswith(fence_char[:3]):
                    break
                code_lines.append(lines[i])
                i += 1
            elements.append({
                "type": _ELEM_CODE_BLOCK,
                "lang": lang,
                "content": '\n'.join(code_lines),
            })
            i += 1
            continue

        # 标题
        heading_match = _RE_HEADING.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": _ELEM_HEADING, "level": level, "text": text})
            i += 1
            continue

        # 表格（连续行，含分隔行）
        if _RE_TABLE_ROW.match(line):
            table_rows = []
            while i < n and _RE_TABLE_ROW.match(lines[i]):
                if not _RE_TABLE_SEP.match(lines[i]):
                    table_rows.append(_parse_table_row(lines[i]))
                i += 1
            if table_rows:
                elements.append({"type": _ELEM_TABLE, "rows": table_rows})
            continue

        # 列表项
        list_match = _RE_LIST_ITEM.match(line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(3)
            elements.append({
                "type": _ELEM_LIST_ITEM,
                "indent": indent,
                "text": text,
            })
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落（合并连续非空行）
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not _RE_HEADING.match(lines[i]) \
                and not _RE_TABLE_ROW.match(lines[i]) and not _RE_LIST_ITEM.match(lines[i]) \
                and not _RE_CODE_FENCE.match(lines[i]):
            para_lines.append(lines[i])
            i += 1
        elements.append({
            "type": _ELEM_PARAGRAPH,
            "text": ' '.join(para_lines),
        })

    return elements


def _parse_table_row(line: str) -> List[str]:
    """解析一行 Markdown 表格，返回单元格内容列表。容错处理。"""
    # 去除首尾管道符
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    cells = [c.strip() for c in line.split('|')]
    return cells


def _extract_markdown_tables(md: str) -> List[List[List[str]]]:
    """提取 Markdown 中所有表格，返回表格列表（每个表格为二维字符串列表）。"""
    lines = md.split('\n')
    tables: List[List[List[str]]] = []
    i = 0
    n = len(lines)

    while i < n:
        if _RE_TABLE_ROW.match(lines[i]):
            table_rows: List[List[str]] = []
            while i < n and _RE_TABLE_ROW.match(lines[i]):
                if not _RE_TABLE_SEP.match(lines[i]):
                    table_rows.append(_parse_table_row(lines[i]))
                i += 1
            if table_rows:
                tables.append(table_rows)
        else:
            i += 1

    return tables


# ======================================================================
# 内部辅助函数 —— Word 写入
# ======================================================================

def _setup_default_font(doc: Document):
    """设置文档默认字体，配置中文字体回退（微软雅黑）。"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)

    # 设置中文字体回退（East Asian）
    rpr = style.element.get_or_add_rPr()
    ea_font = rpr.find(qn('w:rFonts'))
    if ea_font is None:
        from lxml import etree
        ea_font = etree.SubElement(rpr, qn('w:rFonts'))
    ea_font.set(qn('w:eastAsia'), '微软雅黑')


def _set_chinese_font(run, size: Optional[Pt] = None, bold: bool = False):
    """为 Run 对象设置中文字体。"""
    if run is None:
        return
    run.font.name = '微软雅黑'
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    # East Asian 字体
    rpr = run._element.get_or_add_rPr()
    ea_font = rpr.find(qn('w:rFonts'))
    if ea_font is None:
        from lxml import etree
        ea_font = etree.SubElement(rpr, qn('w:rFonts'))
    ea_font.set(qn('w:eastAsia'), '微软雅黑')


def _write_element_to_doc(doc: Document, elem: dict):
    """将解析后的 Markdown 元素写入 Word 文档。"""
    etype = elem["type"]

    if etype == _ELEM_HEADING:
        level = min(elem["level"], 9)  # Word heading 最高支持 9 级
        p = doc.add_heading(elem["text"], level=level)
        for run in p.runs:
            _set_chinese_font(run)

    elif etype == _ELEM_PARAGRAPH:
        p = doc.add_paragraph()
        _add_rich_text(p, elem["text"])

    elif etype == _ELEM_LIST_ITEM:
        indent_level = elem["indent"] // 2
        p = doc.add_paragraph(style='List Bullet')
        # 处理缩进层级
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(indent_level * 0.63)
        _add_rich_text(p, elem["text"])

    elif etype == _ELEM_CODE_BLOCK:
        # 代码块：使用单间距和灰色底纹
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 灰色底纹
        from docx.oxml.ns import qn as _qn
        from lxml import etree
        shd = etree.SubElement(p._element.get_or_add_pPr(), _qn('w:shd'))
        shd.set(_qn('w:fill'), 'F2F2F2')
        shd.set(_qn('w:val'), 'clear')
        run = p.add_run(elem["content"])
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    elif etype == _ELEM_TABLE:
        rows_data = elem["rows"]
        if not rows_data:
            return
        # 计算列数（取最大列数，容错）
        max_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = 'Table Grid'

        for r_idx, row_data in enumerate(rows_data):
            for c_idx in range(max_cols):
                cell = table.cell(r_idx, c_idx)
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                cell.text = cell_text
                # 首行加粗
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            _set_chinese_font(run, size=Pt(10))
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            _set_chinese_font(run, size=Pt(10))


def _add_rich_text(paragraph, text: str):
    """为段落添加富文本（支持 **粗体** 标记）。"""
    parts = _RE_BOLD.split(text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _set_chinese_font(run)
        # 奇数索引为粗体内容
        if idx % 2 == 1:
            run.bold = True


def _add_disclaimer(doc: Document):
    """在文档末尾添加红色 AI 免责声明。"""
    doc.add_paragraph()  # 空行
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_text = (
        "【AI 生成免责声明】本文档由 AI 辅助生成，仅供功能安全分析参考。"
        "所有分析结论和建议措施须由具备资质的功能安全工程师进行人工审查和确认。"
        "AI 生成内容不构成合规性认证的依据，使用者需对最终文档的准确性和完整性负责。"
    )
    run = p.add_run(disclaimer_text)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(9)
    run.italic = True
    _set_chinese_font(run, size=Pt(9))


# ======================================================================
# 内部辅助函数 —— Excel 写入
# ======================================================================

# 样式常量
_HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
_HEADER_FONT = Font(name="微软雅黑", bold=True, color="FFFFFF", size=10)
_NORMAL_FONT = Font(name="微软雅黑", size=10)
_RPN_HIGH_FILL = PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid")  # 红色
_RPN_MED_FILL = PatternFill(start_color="FFE066", end_color="FFE066", fill_type="solid")   # 黄色
_THIN_BORDER = Border(
    left=Side(style='thin'),
    right=Side(style='thin'),
    top=Side(style='thin'),
    bottom=Side(style='thin'),
)

# RPN 列的可能表头名称（中英文）
_RPN_HEADERS = {"rpn", "RPN", "rpn值", "风险优先数"}


def _write_table_to_sheet(ws, rows: List[List[str]]):
    """将二维表格数据写入 Excel Worksheet，并应用 FMEA 格式化。"""
    if not rows:
        return

    header_row = rows[0]
    max_cols = max(len(r) for r in rows)

    # 找到 RPN 所在列
    rpn_col: Optional[int] = None
    for c_idx, cell_val in enumerate(header_row):
        if cell_val.strip() in _RPN_HEADERS:
            rpn_col = c_idx
            break

    # 写入数据
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(max_cols):
            cell_val = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = ws.cell(row=r_idx + 1, column=c_idx + 1, value=cell_val)
            cell.border = _THIN_BORDER
            cell.alignment = Alignment(wrap_text=True, vertical="top")

            if r_idx == 0:
                # 表头样式：蓝底白字
                cell.fill = _HEADER_FILL
                cell.font = _HEADER_FONT
                cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
            else:
                cell.font = _NORMAL_FONT
                # RPN 条件着色
                if rpn_col is not None and c_idx == rpn_col:
                    try:
                        rpn_val = int(cell_val)
                        if rpn_val >= 200:
                            cell.fill = _RPN_HIGH_FILL
                            cell.font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
                        elif rpn_val >= 100:
                            cell.fill = _RPN_MED_FILL
                            cell.font = Font(name="微软雅黑", size=10, bold=True)
                    except (ValueError, TypeError):
                        pass

    # 自动列宽
    _auto_column_width(ws, max_cols, len(rows))

    # 冻结首行
    ws.freeze_panes = "A2"


def _auto_column_width(ws, max_cols: int, max_rows: int):
    """根据内容自动调整列宽。"""
    for c_idx in range(1, max_cols + 1):
        max_width = 8  # 最小宽度
        for r_idx in range(1, min(max_rows + 1, 100)):  # 只采样前100行
            cell = ws.cell(row=r_idx, column=c_idx)
            if cell.value:
                # 粗略估算：中文字符按2个宽度单位
                val_str = str(cell.value)
                width = sum(2 if ord(c) > 127 else 1 for c in val_str[:50])
                max_width = max(max_width, width)
        col_letter = get_column_letter(c_idx)
        ws.column_dimensions[col_letter].width = min(max_width + 2, 50)


def _build_tracking_sheet(ws, tables: List[List[List[str]]]):
    """
    构建「措施跟踪」Sheet，预填从 FMEA 表格中提取的改进措施。
    """
    # 表头
    tracking_headers = [
        "序号", "失效模式ID", "失效模式", "RPN", "AP",
        "现有控制措施", "建议改进措施", "责任人", "计划完成日期",
        "实际完成日期", "状态", "备注"
    ]
    for c_idx, header in enumerate(tracking_headers, start=1):
        cell = ws.cell(row=1, column=c_idx, value=header)
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.border = _THIN_BORDER
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

    # 从 FMEA 表格中提取改进措施行
    row_num = 2
    seq = 1

    # 识别关键列索引
    for table in tables:
        if not table:
            continue
        header = table[0]
        col_map = _build_column_map(header)

        for data_row in table[1:]:
            # 确保数据行足够长
            padded = data_row + [""] * (len(header) - len(data_row))

            fid = _safe_get(padded, col_map.get("id", -1))
            fm = _safe_get(padded, col_map.get("failure_mode", -1))
            rpn = _safe_get(padded, col_map.get("rpn", -1))
            ap = _safe_get(padded, col_map.get("ap", -1))
            existing = _safe_get(padded, col_map.get("existing_ctrl", -1))
            suggestion = _safe_get(padded, col_map.get("suggestion", -1))

            # 只提取有建议措施的条目
            if suggestion or rpn:
                ws.cell(row=row_num, column=1, value=seq).border = _THIN_BORDER
                ws.cell(row=row_num, column=2, value=fid).border = _THIN_BORDER
                ws.cell(row=row_num, column=3, value=fm).border = _THIN_BORDER
                ws.cell(row=row_num, column=4, value=rpn).border = _THIN_BORDER
                ws.cell(row=row_num, column=5, value=ap).border = _THIN_BORDER
                ws.cell(row=row_num, column=6, value=existing).border = _THIN_BORDER
                ws.cell(row=row_num, column=7, value=suggestion).border = _THIN_BORDER
                # 责任人、计划日期、实际日期、状态、备注留空待填
                for c in range(8, 13):
                    ws.cell(row=row_num, column=c, value="").border = _THIN_BORDER

                # 设置字体
                for c in range(1, 13):
                    ws.cell(row=row_num, column=c).font = _NORMAL_FONT
                    ws.cell(row=row_num, column=c).alignment = Alignment(wrap_text=True, vertical="top")

                row_num += 1
                seq += 1

    # 自动列宽
    _auto_column_width(ws, len(tracking_headers), row_num - 1)

    # 冻结首行
    ws.freeze_panes = "A2"


def _build_column_map(header: List[str]) -> dict:
    """
    根据表头名称构建列名到索引的映射，容错处理。
    """
    col_map = {}
    for idx, h in enumerate(header):
        h_lower = h.strip().lower()
        if h_lower in ("失效模式id", "失效模式 id", "fm-id", "fmid", "id"):
            col_map["id"] = idx
        elif h_lower in ("失效模式", "failure mode", "fm"):
            col_map["failure_mode"] = idx
        elif h_lower in ("失效原因", "failure cause", "cause"):
            col_map["cause"] = idx
        elif h_lower in ("失效影响", "failure effect", "effect"):
            col_map["effect"] = idx
        elif h_lower in ("rpn", "rpn值", "风险优先数"):
            col_map["rpn"] = idx
        elif h_lower in ("ap", "行动优先级", "action priority"):
            col_map["ap"] = idx
        elif "现有" in h_lower or "current" in h_lower or "控制" in h_lower:
            col_map["existing_ctrl"] = idx
        elif "建议" in h_lower or "改进" in h_lower or "recommended" in h_lower or "suggested" in h_lower:
            col_map["suggestion"] = idx
        elif h_lower in ("严重度(s)", "严重度", "severity", "s"):
            col_map["severity"] = idx
        elif h_lower in ("发生度(o)", "发生度", "occurrence", "o"):
            col_map["occurrence"] = idx
        elif h_lower in ("检测度(d)", "检测度", "detection", "d"):
            col_map["detection"] = idx
    return col_map


def _safe_get(lst: list, idx: int) -> str:
    """安全获取列表元素，越界返回空字符串。"""
    if idx < 0 or idx >= len(lst):
        return ""
    return lst[idx]
